#Bitlocker Key Finder v3.2
import re
import os
import fnmatch
import shutil
import subprocess
import string
import ctypes
import chardet
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import docx
import pandas as pd
import csv
from striprtf.striprtf import rtf_to_text

pattern = re.compile(r"\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}")
key_pattern = re.compile(r"\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}-\d{6}")
id_pattern = re.compile(r"[a-zA-Z0-9]{8}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]{12}")
data = []
Bit_Keys = []
txt_Files = []
Collected_Keys = []
now = datetime.datetime.now()

# STARTUPINFO to hide the command windows
# startupinfo = subprocess.STARTUPINFO()
# startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
# startupinfo.wShowWindow = subprocess.SW_HIDE

def isAdmin():
    try:
        is_admin = (os.getuid() == 0)
    except AttributeError:
        is_admin = ctypes.windll.shell32.IsUserAnAdmin() != 0
    return is_admin
def parse_docx(gui,file):
    try:
        doc = docx.Document(file)
        for para in doc.paragraphs:
            print(para.text)
            key_match = key_pattern.findall(para.text)
            id_match = id_pattern.findall(para.text)
            if key_match:
                for key in key_match:
                    recovery_id = id_match[0] if id_match else "Unknown"
                    data.append({"File Path": file, "Recovery Key ID": recovery_id, "BitLocker Key": key})
                    gui.log_message(f"File {file} Key ID {recovery_id} Bitlocker Key {key}", "info")
    except Exception as e:
        print(f"Error parsing {file}: {e}")
        return
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    key_match = key_pattern.findall(cell.text)
                    id_match = id_pattern.findall(cell.text)
                    if key_match:
                        for key in key_match:
                            recovery_id = id_match[0] if id_match else "Unknown"
                            data.append({"File Path": file, "Recovery Key ID": recovery_id, "BitLocker Key": key})
    except Exception as e:
        print(f"Error parsing {file}: {e}")
        return

def parse_xlsx(file):
    try:
        doc = pd.read_excel(file)
        for index, row in doc.iterrows():
                # Convert the row to a string and write to the text file
                row_text = ' | '.join([str(item) for item in row.values])  # Join row values with a delimiter
                key_match = key_pattern.findall(row_text)
                id_match = id_pattern.findall(row_text)
                if key_match:
                    for key in key_match:
                        recovery_id = id_match[0] if id_match else "Unknown"
                        data.append({"File Path": file, "Recovery Key ID": recovery_id, "BitLocker Key": key})
    except Exception as e:
        print(f"Error parsing {file}: {e}")
        return
    
def parse_rtf(file):
    try:
        with open(file, 'rb') as f:  # Open in binary mode
            rtf_content = f.read()
            text = rtf_to_text(rtf_content.decode('utf-8'))  # Decode to proper encoding
            key_match = key_pattern.findall(text)
            id_match = id_pattern.findall(text)
            if key_match:
                for key in key_match:
                    recovery_id = id_match[0] if id_match else "Unknown"
                    data.append({"File Path": file, "Recovery Key ID": recovery_id, "BitLocker Key": key})
    except Exception as e:
        print(f"Error parsing {file}: {e}")

def walk(folder):
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith(('.txt', '.TXT', '.bek', '.BEK')):
                txt_Files.append(os.path.join(root, file))
            
def exhaustive_walk(gui, folder):
    for root, dirs, files in os.walk(folder):
        for file in files:
            # if file.endswith(('.txt', '.TXT', '.bek', '.BEK')):
            #     txt_Files.append(os.path.join(root, file))
            if file.endswith(".docx"):
                parse_docx(gui, os.path.join(root, file))
            elif file.endswith(".xlsx"):
                parse_xlsx(gui, os.path.join(root, file))
            elif file.endswith(".rtf"):
                parse_rtf(gui, os.path.join(root, file))


def name_search(gui):
    for ele in txt_Files:
        if fnmatch.fnmatch(ele, "*BitLocker Recovery Key*"):
            Bit_Keys.append(ele)
            gui.log_message(f"Found BitLocker Recovery Key file: {ele}", "success")
        if fnmatch.fnmatch(ele, "*.BEK"):
            Bit_Keys.append(ele)
            gui.log_message(f"Found BEK file: {ele}", "success")

def exhaustive_search(gui):
    for ele in txt_Files:
        try:
            if not os.path.exists(ele):
                gui.log_message(f"File not found: {ele}", "warning")
                continue
            
            too_big = os.path.getsize(ele)
            if too_big >= 1048576:  # Skip files larger than 1MB
                print(f"File too large (>1MB), skipping: {ele}", "warning")
                continue

            # Read the raw contents of the file
            with open(ele, 'rb') as raw_file:
                raw_content = raw_file.read()

            # Detect the encoding
            detected = chardet.detect(raw_content)
            encoding = detected['encoding']

            try:
                decoded_content = raw_content.decode(encoding)
                # k = re.findall(pattern, decoded_content)
                key_match = key_pattern.findall(decoded_content)
                id_match = id_pattern.findall(decoded_content)
                if key_match:
                    for key in key_match:
                        recovery_id = id_match[0] if id_match else "Unknown"
                        data.append({"File Path": ele, "Recovery Key ID": recovery_id, "BitLocker Key": key})
                for key in key_match:
                    Bit_Keys.append(ele)
                    gui.log_message(f"Found BitLocker key in file: {ele}", "success")
                    gui.log_message(f"Key: {key}", "info")
                    Collected_Keys.append(key)
            except UnicodeDecodeError:
                gui.log_message(f"Unable to decode file as {encoding}: {ele}", "warning")
            
        except FileNotFoundError:
            gui.log_message(f"File not found: {ele}", "warning")
        except PermissionError:
            gui.log_message(f"Permission denied: {ele}", "warning")
        except Exception as e:
            print(f"Error processing file {ele}: {str(e)}", "error")

def UTF16LE_string_search(gui):
    for ele in txt_Files:
        try:
            if not os.path.exists(ele):
                gui.log_message(f"File not found: {ele}", "warning")
                continue
            
            too_big = os.path.getsize(ele)
            if too_big >= 1048576:  # Skip files larger than 1MB
                print(f"File too large (>1MB), skipping: {ele}", "warning")
                continue

            # Read the raw contents of the file
            with open(ele, 'rb') as raw_file:
                raw_content = raw_file.read()

            try:
                decoded_content = raw_content.decode('utf-16-le')
                # k = re.findall(pattern, decoded_content)
                key_match = key_pattern.findall(decoded_content)
                id_match = id_pattern.findall(decoded_content)
                if key_match:
                    for key in key_match:
                        recovery_id = id_match[0] if id_match else "Unknown"
                        data.append({"File Path": ele, "Recovery Key ID": recovery_id, "BitLocker Key": key})
                
                        Bit_Keys.append(ele)
                        gui.log_message(f"Found BitLocker key in file: {ele}", "success")
                        gui.log_message(f"Key: {key}", "info")
                        Collected_Keys.append(key)
            except UnicodeDecodeError:
                # gui.log_message(f"Unable to decode file as UTF-16LE: {ele}", "warning")
                pass
            else:
                print(f"File is not UTF-16LE encoded, skipping: {ele}", "info")

        except FileNotFoundError:
            gui.log_message(f"File not found: {ele}", "warning")
        except PermissionError:
            gui.log_message(f"Permission denied: {ele}", "warning")
        except Exception as e:
            gui.log_message(f"Error processing file {ele}: {str(e)}", "error")

def make_key_report(report_folder):
        output_folder = report_folder 
        csv_file = os.path.join(output_folder, "BitlockerKeyReport" + now.strftime("%Y%m%d%H%M%S") + ".csv")
        with open(csv_file, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.DictWriter(file, fieldnames=["File Path", "Recovery Key ID", "BitLocker Key"])
            writer.writeheader()
            writer.writerows(data)
            # print(data)
        
class BitlockerKeyFinderGUI:
    def __init__(self, master):
        self.master = master
        master.title("North Loop Consulting - Bitlocker Key Finder v3.2")
        master.geometry("800x600")
        master.configure(bg="#f0f0f0")

        self.create_widgets()
        self.master.after(100, self.periodic_refresh)  # Start periodic refresh

    def create_widgets(self):
        
        # Find Saved Bitlocker .TXT and .BEK Files
        tk.Label(self.master, text="Find Saved Bitlocker .TXT and .BEK Files", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)

        # Source Directory
        source_frame = tk.Frame(self.master, bg="#f0f0f0")
        source_frame.pack(fill="x", padx=10, pady=5)
        tk.Label(source_frame, text="Select Volume or Directory to Search:", bg="#f0f0f0").pack(side="left")
        self.source_entry = tk.Entry(source_frame, width=50)
        self.source_entry.pack(side="left", expand=True, fill="x", padx=5)
        tk.Button(source_frame, text="Browse", command=self.browse_source, bg="#4CAF50", fg="white", relief=tk.RAISED).pack(side="left")

        # Checkboxes for search options
        checkbox_frame = tk.Frame(self.master, bg="#f0f0f0")
        checkbox_frame.pack(pady=5)
        self.filename_var = tk.BooleanVar()
        tk.Checkbutton(checkbox_frame, text="File Name Search (Fast)", variable=self.filename_var, bg="#f0f0f0").pack(side="left")

        self.utf16le_search_var = tk.BooleanVar()
        tk.Checkbutton(checkbox_frame, text="UTF-16LE String Search", variable=self.utf16le_search_var, bg="#f0f0f0").pack(side="left")

        self.exhaustive_search_var = tk.BooleanVar()
        tk.Checkbutton(checkbox_frame, text="Exhaustive Search - All txt, docx, xlsx, & rtf files (Slow)", variable=self.exhaustive_search_var, bg="#f0f0f0").pack(side="left")

        # Copy Files Option
        self.copy_var = tk.BooleanVar()
        tk.Checkbutton(self.master, text="Copy responsive files to Output Directory", variable=self.copy_var, bg="#f0f0f0").pack(pady=5)

        # Recover Keys from Current Machine
        tk.Label(self.master, text="Recover Keys from Current Machine", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)
        self.mgbde_var = tk.BooleanVar()
        tk.Checkbutton(self.master, text="ADMIN ONLY - Save keys for mounted volumes to Output Directory", variable=self.mgbde_var, bg="#f0f0f0").pack()

        # Output Directory
        output_frame = tk.Frame(self.master, bg="#f0f0f0")
        output_frame.pack(fill="x", padx=10, pady=5)
        tk.Label(output_frame, text="Output Directory:", font=("Arial", 11, "bold"), bg="#f0f0f0").pack(side="left")
        self.output_entry = tk.Entry(output_frame, width=50)
        self.output_entry.pack(side="left", expand=True, fill="x", padx=5)
        self.output_entry.insert(0, os.getcwd())
        tk.Button(output_frame, text="Browse", command=self.browse_output, bg="#4CAF50", fg="white", relief=tk.RAISED).pack(side="left")

        # Buttons
        button_frame = tk.Frame(self.master, bg="#f0f0f0")
        button_frame.pack(pady=10)
        self.find_keys_button = tk.Button(button_frame, text="Find Keys", command=self.start_find_keys_thread, bg="#2196F3", fg="white", relief=tk.RAISED)
        self.find_keys_button.pack(side="left", padx=5)
        
        tk.Button(button_frame, text="Help", command=self.show_help, bg="orange", fg="black", relief=tk.RAISED).pack(side="left", padx=5)
        
        # Console
        console_frame = tk.Frame(self.master, bg="#333333")
        console_frame.pack(fill="both", expand=True, padx=10, pady=5)
        # tk.Label(console_frame, text="Console Output", font=("Arial", 12, "bold"), bg="#333333", fg="white").pack()
        self.console = tk.Text(console_frame, wrap="word", height=15, bg="#ffffff", fg="white")
        self.console.pack(fill="both", expand=True)
        self.console.tag_configure("info", foreground="blue")
        self.console.tag_configure("warning", foreground="orange")
        self.console.tag_configure("error", foreground="red")
        self.console.tag_configure("success", foreground="black")
        self.console.tag_configure("bold", font=("Arial", 10,))
        self.progress = ttk.Progressbar(self.master, orient="horizontal", length=800, mode="determinate")
        self.progress.pack(pady=0)
        clear_frame = tk.Frame(self.master, bg="#f0f0f0")
        clear_frame.pack(fill="x", padx=10, pady=(0, 10))  # Add padding at the bottom
        tk.Button(clear_frame, text="Clear Window", command=self.clear_console, bg="#FF5722", fg="white", relief=tk.RAISED).pack()


    def browse_source(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.source_entry.delete(0, tk.END)
            self.source_entry.insert(0, folder_path)

    def browse_output(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, folder_path)

    def show_help(self):
        help_message = (
            "Copyright 2024 North Loop Consulting\n"
            "Bitlocker Key Finder\n\n"
            "1. Select the directory to search for Bitlocker Recovery Keys or BEK files.\n"
            "2. Choose search options:\n"
            "   - File Name Search: A quick search for file names consistent with key files.\n"
            "   - UTF-16LE String Search: Searches for Bitlocker keys in UTF-16LE encoded files.\n"
            "   - Exhaustive String Search: Performs a search of all .txt files smaller than 1MB for keys.\n"
            "3. Optionally, enable the Copy Files option to copy found files to the output directory.\n"
            "4. Optionally, enable the recovery of keys from the current machine (ADMIN ONLY).\n"
            "5. Choose the output directory to save results.\n"
            "6. Click 'Find Keys' to start the search."
        )
        messagebox.showinfo("Help", help_message)

    def start_find_keys_thread(self):
        # Disable the Find Keys button
        self.find_keys_button.config(state=tk.DISABLED)
        # Start the find keys process in a separate thread
        thread = threading.Thread(target=self.find_keys)
        thread.start()
    
    

    def find_keys(self):
        global Bit_Keys, txt_Files
        Bit_Keys = []
        txt_Files = []
        folder_path = self.source_entry.get()
        
        # Traverse the directory and find .txt and .BEK files
        walk(folder_path)
        total_files = len(txt_Files)
        self.progress['value'] = 0
        self.progress['maximum'] = total_files

        if self.filename_var.get():
            self.log_message(f"Searching for file names in {folder_path}", "info")
            name_search(self)

        if self.exhaustive_search_var.get():
            self.log_message(f"Conducting exhaustive string search in {folder_path}", "info")
            exhaustive_walk(self, folder_path)
            exhaustive_search(self)
            make_key_report(self.output_entry.get())
            
        

        if self.utf16le_search_var.get():
            self.log_message(f"Conducting UTF-16LE string search in {folder_path}", "info")
            UTF16LE_string_search(self)

        self.log_message(f"Total BitLocker keys/files found: {len(Bit_Keys)}", "success")
        
        if self.copy_var.get():
            self.copy_key_files()
        
        if self.mgbde_var.get():
            self.get_active_keys()
        
        self.log_message("SEARCH COMPLETE", "success")
        
        # Re-enable the Find Keys button
        self.find_keys_button.config(state=tk.NORMAL)

    


    def copy_key_files(self):
        output_folder = self.output_entry.get()
        if not os.path.isdir(output_folder):
            self.log_message("Invalid output directory. Please select a valid directory.", "warning")
            return
        for file in Bit_Keys:
            try:
                shutil.copy(file, output_folder)
                self.log_message(f"Copied file: {file}", "success")
            except Exception as e:
                self.log_message(f"Error copying file {file}: {str(e)}", "error")

    def get_active_keys(self):
        if not isAdmin():
            self.log_message("Admin rights are required to retrieve BitLocker keys.", "warning")
            return
        output_folder = self.output_entry.get()
        comp_name = os.environ['COMPUTERNAME']  #gets target computer name for report title
        comp_name = comp_name.strip('\\')
        key_report = os.path.join(output_folder, comp_name + '-BitlockerReport.txt')
        Drive_letters = ['%s:' % d for d in string.ascii_uppercase if os.path.exists('%s:' % d)]   #Produces list of volumes on target system

        
        if not os.path.isdir(output_folder):
            self.log_message("Invalid output directory. Please select a valid directory.", "warning")
            return
        with open(key_report, 'w') as report:
            report.write("Bitlocker Key Finder v3.0 \n")  #writing the header for the report 1) Version 2) Date 3)User of System
            report.write(now.strftime("%Y-%m-%d, %H:%M:%S")) 
            report.write("\nUser Account Used: ")
            report.write(os.getlogin())
            report.write("\n\n")
        try:
            # volumes = subprocess.check_output(["manage-bde", "-status"], startupinfo=startupinfo).decode("utf-8")
            volumes = subprocess.check_output(["manage-bde", "-status"]).decode("utf-8")

            self.log_message(volumes, "info")
            volume_lines = volumes.splitlines()
            with open(key_report, "a") as key_file:
                for line in volume_lines:
                    
                    if "Volume " in line:
                        volume = line.split()[1]
                        print(volume)
                        try:
                            recovery_keys = subprocess.check_output(["manage-bde", "-protectors", "-get", volume], startupinfo=startupinfo).decode("utf-8")
                            key_file.write(f"Bitlocker key found for {volume}!\n\n")
                            key_file.write(recovery_keys)
                            self.log_message(f"BitLocker key for volume {volume} written to report at {key_report}", "success")
                            # self.log_message(f"{recovery_keys}", "info")
                        except subprocess.CalledProcessError:
                            # self.log_message(f"No BitLocker credentials found for {volume}", "warning")
                            key_file.write(f"No BitLocker credentials found for {volume}\n\n")
        except Exception as e:
            self.log_message(f"Error retrieving BitLocker keys: {str(e)}", "error")
        

    def log_message(self, message, level="info"):
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        formatted_message = f"[{timestamp}] {message}\n"
        self.console.insert(tk.END, formatted_message, (level, "bold"))
        self.console.see(tk.END)
        self.master.after(100, lambda: self.master.update_idletasks())  # Regular GUI update

    def clear_console(self):
        self.console.delete(1.0, tk.END)

    def periodic_refresh(self):
        self.master.update_idletasks()
        self.master.after(100, self.periodic_refresh)

# Create and run the Tkinter application
root = tk.Tk()
app = BitlockerKeyFinderGUI(root)
root.mainloop()
